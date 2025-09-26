import fitz  # PyMuPDF
from docx import Document

def extraer_texto(origen):
    """
    Extrae texto desde una ruta local (str) o desde un archivo subido (FileStorage).
    """
    # 📁 Si es una ruta local
    if isinstance(origen, str):
        if origen.endswith(".pdf"):
            doc = fitz.open(origen)
            try:
                return "\n".join([pagina.get_text() for pagina in doc])
            finally:
                doc.close()
        elif origen.endswith(".docx"):
            doc = Document(origen)
            # python-docx no requiere close explícito
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise ValueError("Formato no soportado")

    # 🌐 Si es un archivo subido desde Flask
    elif hasattr(origen, "filename"):
        nombre = origen.filename
        if nombre.endswith(".pdf"):
            origen.seek(0)
            data = origen.read()
            doc = fitz.open(stream=data, filetype="pdf")
            try:
                return "\n".join([pagina.get_text() for pagina in doc])
            finally:
                doc.close()
        elif nombre.endswith(".docx"):
            origen.seek(0)
            doc = Document(origen)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise ValueError("Formato no soportado")

    else:
        raise TypeError("Origen debe ser una ruta (str) o un archivo subido (FileStorage)")
